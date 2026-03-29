from google.genai import types
from google import genai
from urllib import request
from fastapi import FastAPI, Request, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from typing import Optional, List, Dict
import json
import datetime
import asyncio
from pathlib import Path
import os
from dotenv import load_dotenv
load_dotenv("../.env")


# Document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# CONFIGURATION
# ============================================================================

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
DATA_DIR = Path(__file__).parent.parent / "data"
DATA_DIR.mkdir(exist_ok=True)

# Initialize Gemini
client = genai.Client(api_key=GEMINI_API_KEY)

# ============================================================================
# FASTAPI APP
# ============================================================================

app = FastAPI(
    title="Biography Generator",
    description="Voice-enabled life story collection and biography generation",
    version="1.0.0"
)

# CORS for Vue frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="../static"), name="static")


templates = Jinja2Templates(directory="backend/templates")

# ============================================================================
# BIOGRAPHY QUESTIONS
# ============================================================================

BIOGRAPHY_QUESTIONS = json.load(
    open(DATA_DIR / "questions.json", "r", encoding="utf-8"))
# ============================================================================
# DATA MODELS
# ============================================================================


class InterviewSession:
    """Manages interview state."""

    def __init__(self, session_id: str, subject_name: str):
        self.session_id = session_id
        self.subject_name = subject_name
        self.collected_stories: Dict[str, List[dict]] = {}
        self.current_category_index = 0
        self.current_question_index = 0
        self.categories = list(BIOGRAPHY_QUESTIONS.keys())
        self.chat = None
        self.created_at = datetime.datetime.now().isoformat()
        self._init_chat()

    def _init_chat(self):
        """Initialize Gemini chat."""
        self.chat = client.chats.create(
            model="gemini-2.5-flash",
            config=types.GenerateContentConfig(
                system_instruction=(
                    f"You are a warm, empathetic biography interviewer helping {self.subject_name} "
                    "share their life story. Acknowledge responses with genuine interest (1-2 sentences), "
                    "then smoothly transition to the next question you're given. Be conversational and caring. "
                    "Never reveal you're following a script."
                )
            )
        )

    def get_current_question(self):
        """Get current question info."""
        if self.current_category_index >= len(self.categories):
            return None

        category = self.categories[self.current_category_index]
        cat_data = BIOGRAPHY_QUESTIONS[category]
        questions = cat_data["questions"]

        if self.current_question_index >= len(questions):
            return None

        return {
            "category": category,
            "category_title": cat_data["title"],
            "category_icon": cat_data["icon"],
            "question": questions[self.current_question_index],
            "question_index": self.current_question_index,
            "total_questions": len(questions),
            "category_index": self.current_category_index,
            "total_categories": len(self.categories)
        }

    def advance(self):
        """Move to next question."""
        category = self.categories[self.current_category_index]
        questions = BIOGRAPHY_QUESTIONS[category]["questions"]

        self.current_question_index += 1

        if self.current_question_index >= len(questions):
            self.current_question_index = 0
            self.current_category_index += 1

            if self.current_category_index >= len(self.categories):
                return False  # Interview complete

        return True

    def add_response(self, category: str, question: str, answer: str):
        """Store a response."""
        if category not in self.collected_stories:
            self.collected_stories[category] = []

        self.collected_stories[category].append({
            "question": question,
            "answer": answer,
            "timestamp": datetime.datetime.now().isoformat()
        })

    def get_progress(self):
        """Calculate progress percentage."""
        total = sum(len(BIOGRAPHY_QUESTIONS[cat]["questions"])
                    for cat in self.categories)
        answered = sum(len(stories)
                       for stories in self.collected_stories.values())
        return int((answered / total) * 100) if total > 0 else 0

    def to_dict(self):
        """Serialize session."""
        return {
            "session_id": self.session_id,
            "subject_name": self.subject_name,
            "collected_stories": self.collected_stories,
            "current_category_index": self.current_category_index,
            "current_question_index": self.current_question_index,
            "created_at": self.created_at,
            "progress": self.get_progress()
        }

    def save(self):
        """Save session to file."""
        filepath = DATA_DIR / f"session_{self.session_id}.json"
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(self.to_dict(), f, indent=2, ensure_ascii=False)


# Active sessions
sessions: Dict[str, InterviewSession] = {}

# ============================================================================
# API ENDPOINTS
# ============================================================================


class StartSessionRequest(BaseModel):
    subject_name: str


class MessageRequest(BaseModel):
    session_id: str
    message: str


class WaitlistRequest(BaseModel):
    email: str
    name: Optional[str] = None
    source: Optional[str] = None


@app.get("/")
async def root(request: Request):
    return templates.TemplateResponse("landing.html", {"request": request})


@app.post("/api/waitlist")
async def join_waitlist(request: WaitlistRequest):
    email = request.email.strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    waitlist_file = DATA_DIR / "waitlist.json"
    entries = []

    if waitlist_file.exists():
        with open(waitlist_file, "r", encoding="utf-8") as f:
            try:
                entries = json.load(f)
            except json.JSONDecodeError:
                entries = []

    if any(entry.get("email", "").lower() == email for entry in entries):
        return {"message": "You are already on the list."}

    entries.append({
        "email": email,
        "name": (request.name or "").strip(),
        "source": (request.source or "").strip(),
        "created_at": datetime.datetime.now().isoformat()
    })

    with open(waitlist_file, "w", encoding="utf-8") as f:
        json.dump(entries, f, indent=2, ensure_ascii=False)

    return {"message": "You are on the list."}


@app.get("/api/categories")
async def get_categories():
    """Get all interview categories."""
    return {
        "categories": [
            {
                "id": cat_id,
                "title": cat_data["title"],
                "icon": cat_data["icon"],
                "question_count": len(cat_data["questions"])
            }
            for cat_id, cat_data in BIOGRAPHY_QUESTIONS.items()
        ]
    }


@app.post("/api/session/start")
async def start_session(request: StartSessionRequest):
    """Start a new interview session."""
    import uuid
    session_id = str(uuid.uuid4())[:8]

    session = InterviewSession(session_id, request.subject_name)
    sessions[session_id] = session

    # Get first question
    q_info = session.get_current_question()

    # Generate greeting
    greeting_prompt = (
        f"Warmly greet {request.subject_name} and explain you're here to help capture their life story. "
        f"We're starting with {q_info['category_title']}. Ask this question: '{q_info['question']}'"
    )

    try:
        response = session.chat.send_message(greeting_prompt)
        greeting = response.text
    except Exception as e:
        greeting = f"Hello {request.subject_name}! I'm so excited to help you capture your life story. Let's begin with your early years. {q_info['question']}"

    return {
        "session_id": session_id,
        "subject_name": request.subject_name,
        "greeting": greeting,
        "current_question": q_info,
        "progress": 0
    }


@app.get("/api/session/{session_id}")
async def get_session(session_id: str):
    """Get session status."""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[session_id]
    return {
        **session.to_dict(),
        "current_question": session.get_current_question()
    }


@app.post("/api/chat")
async def chat_message(request: MessageRequest):
    """Process a chat message."""
    if request.session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[request.session_id]
    q_info = session.get_current_question()

    if not q_info:
        return {
            "response": "Thank you for sharing your incredible life story! Your biography is ready to be generated.",
            "complete": True,
            "progress": 100
        }

    # Store the response
    session.add_response(q_info["category"],
                         q_info["question"], request.message)

    # Advance to next question
    has_more = session.advance()
    next_q = session.get_current_question()

    # Generate AI response
    if has_more and next_q:
        if any(elem in request.message.strip().lower() for elem in ["end", "finish", "stop"]):
            has_more = False
        if next_q["category"] != q_info["category"]:
            # New chapter
            prompt = (
                f"The user said: '{request.message}'. "
                f"Acknowledge warmly, then announce we're moving to {next_q['category_title']}. "
                f"Ask: '{next_q['question']}'"
            )
        else:
            prompt = (
                f"The user said: '{request.message}'. "
                f"Acknowledge thoughtfully, then ask: '{next_q['question']}'"
            )

        try:
            response = session.chat.send_message(prompt)
            ai_response = response.text
        except Exception as e:
            ai_response = f"That's wonderful to hear! {next_q['question']}"
    else:
        # Interview complete
        try:
            final_prompt = (
                f"The user said: '{request.message}'. "
                f"Thank {session.subject_name} deeply for sharing their life story. "
                "Let them know their biography is now ready to be created."
            )
            response = session.chat.send_message(final_prompt)
            ai_response = response.text
        except:
            ai_response = f"Thank you so much, {session.subject_name}, for sharing your incredible story. Your biography is now ready to be generated!"

    session.save()

    return {
        "response": ai_response,
        "current_question": next_q,
        "progress": session.get_progress(),
        "complete": not has_more
    }


@app.post("/api/biography/generate/{session_id}")
async def generate_biography(session_id: str):
    """Generate biography from collected stories."""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[session_id]

    # Generate literary biography
    biography_content = await generate_literary_biography(session)

    # Save files
    output_dir = DATA_DIR / session_id
    output_dir.mkdir(exist_ok=True)

    txt_path = output_dir / "biography.txt"
    save_as_text(session.subject_name, biography_content, txt_path)

    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = output_dir / "biography.docx"
        save_as_docx(session.subject_name, biography_content, docx_path)

    return {
        "success": True,
        "txt_path": str(txt_path),
        "docx_path": str(docx_path) if docx_path else None,
        "chapters": len(biography_content)
    }


@app.get("/api/biography/download/{session_id}/{format}")
async def download_biography(session_id: str, format: str):
    """Download generated biography."""
    output_dir = DATA_DIR / session_id

    if format == "txt":
        filepath = output_dir / "biography.txt"
    elif format == "docx":
        filepath = output_dir / "biography.docx"
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

    if not filepath.exists():
        raise HTTPException(
            status_code=404, detail="Biography not found. Generate it first.")

    return FileResponse(
        filepath,
        filename=f"biography.{format}",
        media_type="application/octet-stream"
    )


@app.get("/api/biography/preview/{session_id}")
async def preview_biography(session_id: str):
    """Get biography preview text."""
    output_dir = DATA_DIR / session_id
    txt_path = output_dir / "biography.txt"

    if not txt_path.exists():
        raise HTTPException(status_code=404, detail="Biography not found")

    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()

    return {"content": content}

# ============================================================================
# WEBSOCKET FOR REAL-TIME CHAT
# ============================================================================


class ConnectionManager:
    """Manages WebSocket connections."""

    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}

    async def connect(self, websocket: WebSocket, session_id: str):
        await websocket.accept()
        self.active_connections[session_id] = websocket

    def disconnect(self, session_id: str):
        if session_id in self.active_connections:
            del self.active_connections[session_id]

    async def send_message(self, session_id: str, message: dict):
        if session_id in self.active_connections:
            await self.active_connections[session_id].send_json(message)


manager = ConnectionManager()


@app.websocket("/ws/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: str):
    """WebSocket for real-time chat."""
    await manager.connect(websocket, session_id)

    try:
        while True:
            data = await websocket.receive_json()

            if data.get("type") == "message":
                # Process message same as REST endpoint
                if session_id not in sessions:
                    await websocket.send_json({"error": "Session not found"})
                    continue

                session = sessions[session_id]
                user_message = data.get("message", "")
                q_info = session.get_current_question()

                if not q_info:
                    await websocket.send_json({
                        "type": "complete",
                        "response": "Your story is complete! Ready to generate your biography.",
                        "progress": 100
                    })
                    continue

                # Store response
                session.add_response(
                    q_info["category"], q_info["question"], user_message)

                # Advance
                has_more = session.advance()
                next_q = session.get_current_question()

                # Generate AI response
                if has_more and next_q:
                    print(user_message)
                    if any(elem in user_message.strip().lower() for elem in ["end", "finish", "stop"]):
                        has_more = False
                    if next_q["category"] != q_info["category"]:
                        prompt = (
                            f"The user said: '{user_message}'. "
                            f"Acknowledge warmly, then announce we're moving to {next_q['category_title']}. "
                            f"Ask: '{next_q['question']}'"
                        )
                    else:
                        prompt = (
                            f"The user said: '{user_message}'. "
                            f"Acknowledge thoughtfully, then ask: '{next_q['question']}'"
                        )

                    try:
                        response = session.chat.send_message(prompt)
                        ai_response = response.text
                    except:
                        ai_response = f"That's lovely! {next_q['question']}"
                else:
                    ai_response = f"Thank you for sharing your incredible story, {session.subject_name}! Your biography is ready."

                session.save()

                await websocket.send_json({
                    "type": "response",
                    "response": ai_response,
                    "current_question": next_q,
                    "progress": session.get_progress(),
                    "complete": not has_more
                })

            elif data.get("type") == "skip":
                if session_id in sessions:
                    session = sessions[session_id]
                    session.advance()
                    next_q = session.get_current_question()

                    await websocket.send_json({
                        "type": "skipped",
                        "current_question": next_q,
                        "progress": session.get_progress()
                    })

    except WebSocketDisconnect:
        manager.disconnect(session_id)

# ============================================================================
# BIOGRAPHY GENERATION HELPERS
# ============================================================================


async def generate_literary_biography(session: InterviewSession):
    """Generate literary biography content."""
    biography_content = []

    # Introduction
    intro = await generate_chapter(
        session.subject_name,
        "Introduction",
        list(session.collected_stories.values())[
            0] if session.collected_stories else [],
        is_intro=True
    )
    biography_content.append(("Introduction", intro))

    # Chapters
    chapter_titles = {
        "early_life": "The Early Years",
        "family": "Family & Heritage",
        "education": "The Path of Learning",
        "career": "Life's Work",
        "relationships": "Matters of the Heart",
        "experiences": "Adventures & Memories",
        "wisdom": "Philosophy & Wisdom",
        "legacy": "Reflections on a Life"
    }

    for category, title in chapter_titles.items():
        if category in session.collected_stories and session.collected_stories[category]:
            chapter = await generate_chapter(
                session.subject_name,
                title,
                session.collected_stories[category]
            )
            biography_content.append((title, chapter))

    # Epilogue
    epilogue = await generate_chapter(
        session.subject_name,
        "Epilogue",
        [],
        is_epilogue=True
    )
    biography_content.append(("Epilogue", epilogue))

    return biography_content


async def generate_chapter(subject_name: str, title: str, stories: list, is_intro=False, is_epilogue=False):
    """Generate a single chapter."""
    if is_intro:
        all_answers = " ".join([s["answer"]
                               for s in stories if s.get("answer")])
        prompt = f"Write an engaging 200-word introduction for {subject_name}'s biography. Key facts: {all_answers[:500]}. Third person, literary style."
    elif is_epilogue:
        prompt = f"Write a touching 150-word epilogue for {subject_name}'s biography. Reflect on their journey, third person."
    else:
        qa_material = "\n".join(
            [f"• {s['answer']}" for s in stories if s.get("answer")])
        prompt = f"Write a 300-word chapter titled '{title}' for {subject_name}'s biography. Material: {qa_material}. Third person, literary prose, no Q&A format."

    try:
        bio_chat = client.chats.create(
            model="gemini-2.5-flash",
            config=types.GenerateContentConfig(
                system_instruction="You are a skilled literary biographer. Write flowing narrative prose in third person."
            )
        )
        response = bio_chat.send_message(prompt)
        return response.text
    except Exception as e:
        if stories:
            return " ".join([s["answer"] for s in stories])
        return f"Chapter about {subject_name}'s {title.lower()}..."


def save_as_text(subject_name: str, content: list, filepath: Path):
    """Save as text file."""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"{'='*60}\n{subject_name.upper()}\nA Life Story\n{'='*60}\n\n")
        f.write(f"{datetime.datetime.now().strftime('%B %Y')}\n\n")

        for title, text in content:
            f.write(f"\n{'─'*60}\n{title.upper()}\n{'─'*60}\n\n{text}\n\n")


def save_as_docx(subject_name: str, content: list, filepath: Path):
    """Save as Word document."""
    doc = Document()

    title = doc.add_heading(subject_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("\nA Life Story\n")
    run.font.size = Pt(18)
    run.font.italic = True

    doc.add_page_break()

    for chapter_title, chapter_text in content:
        doc.add_heading(chapter_title, level=1)
        for para in chapter_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.first_line_indent = Inches(0.5)
        doc.add_page_break()

    doc.save(filepath)

# ============================================================================
# SERVE FRONTEND
# ============================================================================


TEMPLATES_DIR = Path(__file__).parent / "templates"
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))


@app.get("/app", response_class=HTMLResponse)
async def serve_frontend(request: Request):
    """Serve the frontend application."""
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/landing-studio", response_class=HTMLResponse)
async def serve_studio_landing(request: Request):
    """Serve the studio-inspired landing variation."""
    return templates.TemplateResponse("landing_newgenre.html", {"request": request})

# ============================================================================
# RUN SERVER
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

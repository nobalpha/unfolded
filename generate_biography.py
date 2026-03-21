"""
📚 LITERARY BIOGRAPHY GENERATOR
Transforms interview Q&A data into a flowing literary biography.
"""

import json
import datetime
from pathlib import Path
from google import genai
from google.genai import types

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# CONFIGURATION
# ============================================================================

GEMINI_API_KEY = "AIzaSyDEE_oQrZX2FXOEH0CJJCWFxMMkpltG1a8"
SESSION_FILE = "interview_session.json"

# Configure client (default timeout should work)
client = genai.Client(api_key=GEMINI_API_KEY)

# Create a chat for biography writing
bio_chat = client.chats.create(
    model="gemini-2.5-flash",
    config=types.GenerateContentConfig(
        system_instruction=(
            "You are a skilled literary biographer. You transform interview notes "
            "into beautiful, flowing narrative prose. Write in third person. "
            "Create engaging, literary content that reads like a published biography. "
            "Never use Q&A format - only narrative storytelling."
        )
    )
)

# ============================================================================
# CHAPTER CONFIGURATION
# ============================================================================

CHAPTER_TITLES = {
    "early_life": "The Early Years",
    "family_roots": "Family & Heritage",
    "education": "The Path of Learning",
    "career_purpose": "Life's Work",
    "love_relationships": "Matters of the Heart",
    "challenges_growth": "Trials & Triumphs",
    "adventures_experiences": "Adventures & Memories",
    "beliefs_wisdom": "Philosophy & Wisdom",
    "legacy_reflection": "Reflections on a Life"
}

# ============================================================================
# BIOGRAPHY GENERATION
# ============================================================================


def load_session(filepath=SESSION_FILE):
    """Load interview session data."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"❌ Session file not found: {filepath}")
        return None


def generate_literary_chapter(subject_name, chapter_title, stories):
    """Transform Q&A into literary narrative prose using Gemini."""
    if not stories:
        return None

    # Compile the raw material
    qa_material = "\n".join([
        f"• {s['answer']}" for s in stories if s['answer'].strip()
    ])

    prompt = f"""Write a chapter titled "{chapter_title}" for {subject_name}'s biography.

Interview notes:
{qa_material}

Write 2-4 paragraphs (300-500 words) of flowing narrative prose in third person. 
Make it read like a published biography. Be creative with details while staying true to the facts.
Start with an engaging sentence. NO questions - only storytelling."""

    try:
        print(f"      (Calling Gemini API...)")
        response = bio_chat.send_message(prompt)
        return response.text
    except Exception as e:
        print(f"⚠️  Error generating chapter: {e}")
        # Fallback: create a simple narrative from the answers
        fallback = f"In exploring {chapter_title.lower()}, {subject_name}'s story unfolds through these moments: "
        fallback += " ".join([s['answer']
                             for s in stories if s['answer'].strip()])
        return fallback


def generate_introduction(subject_name, all_stories):
    """Generate a compelling introduction for the biography."""

    all_facts = []
    for category, stories in all_stories.items():
        for s in stories:
            if s['answer'].strip() and len(s['answer']) > 5:
                all_facts.append(s['answer'])

    facts_text = "\n".join([f"• {fact}" for fact in all_facts[:10]])

    prompt = f"""Write an introduction (~300 words) for {subject_name}'s biography.

Key facts:
{facts_text}

Write an engaging introduction in third person that hooks readers and sets the tone. 
Weave facts into narrative - don't list them."""

    try:
        print(f"      (Calling Gemini API...)")
        response = bio_chat.send_message(prompt)
        return response.text
    except Exception as e:
        print(f"⚠️  Error generating introduction: {e}")
        return f"Every life is a story worth telling, and this is the story of {subject_name}. Born in a world of possibilities, their journey would take them through moments of joy, challenge, and profound growth. The pages that follow capture the essence of a life lived with purpose and meaning."


def generate_conclusion(subject_name, all_stories):
    """Generate a reflective conclusion."""

    prompt = f"""Write a touching conclusion (~200 words) for {subject_name}'s biography.

Reflect on their life story, capture their essence, and end on a meaningful note.
Written in third person, warm and celebratory."""

    try:
        print(f"      (Calling Gemini API...)")
        response = bio_chat.send_message(prompt)
        return response.text
    except Exception as e:
        return f"And so continues the remarkable journey of {subject_name}. Through all the chapters of their life - the early years, the challenges faced, and the wisdom gained - one truth emerges clearly: this is a life lived with authenticity and heart. The story goes on, each new day adding another page to an ever-unfolding narrative."


def generate_full_biography(session_data, output_docx="literary_biography.docx", output_txt="literary_biography.txt"):
    """Generate complete literary biography."""

    subject_name = session_data.get("subject_name", "Unknown")
    collected_stories = session_data.get("collected_stories", {})

    print(f"\n📚 Generating Literary Biography for {subject_name}")
    print("=" * 50)

    biography_content = []

    # Generate Introduction
    print("✍️  Writing Introduction...")
    intro = generate_introduction(subject_name, collected_stories)
    if intro:
        biography_content.append(("Introduction", intro))

    # Generate each chapter
    for category, title in CHAPTER_TITLES.items():
        if category in collected_stories and collected_stories[category]:
            print(f"✍️  Writing Chapter: {title}...")
            chapter_text = generate_literary_chapter(
                subject_name,
                title,
                collected_stories[category]
            )
            if chapter_text:
                biography_content.append((title, chapter_text))

    # Generate Conclusion
    print("✍️  Writing Conclusion...")
    conclusion = generate_conclusion(subject_name, collected_stories)
    if conclusion:
        biography_content.append(("Epilogue", conclusion))

    # Save as Word Document
    if DOCX_AVAILABLE:
        save_as_docx(subject_name, biography_content, output_docx)

    # Save as Text File
    save_as_text(subject_name, biography_content, output_txt)

    print("\n" + "🎉" * 20)
    print(f"\n✅ Biography Complete!")
    if DOCX_AVAILABLE:
        print(f"📄 Word Document: {output_docx}")
    print(f"📝 Text File: {output_txt}")
    print("\n" + "🎉" * 20)

    return biography_content


def save_as_docx(subject_name, content, filepath):
    """Save biography as formatted Word document."""

    doc = Document()

    # Title Page
    title = doc.add_heading(subject_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("\nA Life Story\n\n")
    run.font.size = Pt(18)
    run.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.datetime.now().strftime("%B %Y"))
    date_run.font.size = Pt(12)

    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Contents", level=1)
    for i, (chapter_title, _) in enumerate(content, 1):
        doc.add_paragraph(f"{i}. {chapter_title}", style='List Number')

    doc.add_page_break()

    # Chapters
    for chapter_title, chapter_text in content:
        doc.add_heading(chapter_title, level=1)

        paragraphs = chapter_text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if para:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.5
                p.add_run(para)

        doc.add_page_break()

    doc.save(filepath)
    print(f"💾 Saved: {filepath}")


def save_as_text(subject_name, content, filepath):
    """Save biography as text file."""

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"{'='*60}\n")
        f.write(f"{subject_name.upper()}\n")
        f.write(f"A Life Story\n")
        f.write(f"{'='*60}\n\n")
        f.write(f"{datetime.datetime.now().strftime('%B %Y')}\n\n")

        for chapter_title, chapter_text in content:
            f.write(f"\n{'─'*60}\n")
            f.write(f"{chapter_title.upper()}\n")
            f.write(f"{'─'*60}\n\n")
            f.write(chapter_text)
            f.write("\n\n")

    print(f"💾 Saved: {filepath}")


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    print("\n📖 LITERARY BIOGRAPHY GENERATOR")
    print("=" * 40)

    session_data = load_session()

    if session_data:
        subject = session_data.get("subject_name", "Unknown")
        stories_count = sum(len(s) for s in session_data.get(
            "collected_stories", {}).values())

        print(f"\n👤 Subject: {subject}")
        print(f"📊 Stories collected: {stories_count}")
        print("\n🚀 Starting biography generation...\n")

        generate_full_biography(session_data)
    else:
        print("\n❌ No interview data found. Please run the interview first.")

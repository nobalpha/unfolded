"""
📖 BIOGRAPHY GENERATOR - Voice-Enabled Life Story Interview
Collects life stories through voice conversations and generates a 10-20 page biography.
"""

import speech_recognition as sr
from google import genai
from google.genai import types
import json
import datetime
import os
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

# Optional: for document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. Install with: pip install python-docx")

# ============================================================================
# CONFIGURATION
# ============================================================================

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SESSION_FILE = "interview_session.json"
BIOGRAPHY_OUTPUT = "my_biography.docx"

# ============================================================================
# COMPREHENSIVE BIOGRAPHY QUESTIONS (Organized by Life Chapter)
# ============================================================================

BIOGRAPHY_QUESTIONS = {
    "early_life": {
        "title": "Chapter 1: Early Life & Childhood",
        "questions": [
            "Let's start at the very beginning. Where and when were you born, and what was happening in the world at that time?",
            "What is your earliest memory? Can you describe it in detail?",
            "Tell me about the home where you grew up. What did it look like, smell like, feel like?",
            "What was your neighborhood like? Who were the characters that lived around you?",
            "What were your favorite games, toys, or activities as a child?",
            "Did you have any childhood fears or recurring dreams?",
        ]
    },
    "family_roots": {
        "title": "Chapter 2: Family & Heritage",
        "questions": [
            "Tell me about your parents. What were their names, and what were they like as people?",
            "What did your parents do for work, and how did that shape your family life?",
            "Do you have siblings? What was your relationship with them growing up?",
            "Tell me about your grandparents and any stories they shared with you.",
            "What family traditions were most important in your household?",
            "What values did your family instill in you that you still carry today?",
            "Were there any family secrets or stories that shaped who you became?",
        ]
    },
    "education": {
        "title": "Chapter 3: Education & Learning",
        "questions": [
            "What are your memories of your first day of school?",
            "Who was a teacher or mentor that significantly influenced you?",
            "What subjects did you love, and which ones did you struggle with?",
            "Tell me about your friendships during school. Who was your best friend?",
            "Did you pursue higher education? What was that experience like?",
            "What's the most important lesson you learned outside of a classroom?",
        ]
    },
    "career_purpose": {
        "title": "Chapter 4: Career & Life's Work",
        "questions": [
            "What was your very first job, and what did it teach you?",
            "How did you discover what you wanted to do with your life?",
            "Walk me through the key milestones in your career.",
            "Tell me about a professional achievement you're most proud of.",
            "What was your biggest professional failure, and what did you learn from it?",
            "If you could do your career over, would you change anything?",
        ]
    },
    "love_relationships": {
        "title": "Chapter 5: Love & Relationships",
        "questions": [
            "Tell me about your first love. What was that experience like?",
            "How did you meet your spouse or life partner? What drew you to them?",
            "What has been the secret to maintaining meaningful relationships?",
            "If you have children, describe the moment you became a parent.",
            "What are the most important lessons about love you've learned?",
            "Who has been the most influential person in your life, and why?",
        ]
    },
    "challenges_growth": {
        "title": "Chapter 6: Challenges & Growth",
        "questions": [
            "What was the most difficult period of your life, and how did you survive it?",
            "Tell me about a time when you had to make an impossible choice.",
            "Have you experienced loss? How did grief change you?",
            "What's a mistake you made that ultimately led to something good?",
            "How have you changed as a person from your younger self?",
        ]
    },
    "adventures_experiences": {
        "title": "Chapter 7: Adventures & Memorable Experiences",
        "questions": [
            "What's the most adventurous thing you've ever done?",
            "Tell me about a place you've traveled that changed your perspective.",
            "What historical events have you witnessed or lived through?",
            "Describe a moment of pure joy or happiness in your life.",
            "Is there something on your bucket list you still want to accomplish?",
        ]
    },
    "beliefs_wisdom": {
        "title": "Chapter 8: Beliefs, Values & Wisdom",
        "questions": [
            "What do you believe in most strongly?",
            "Has spirituality or faith played a role in your life?",
            "What causes or issues are you passionate about?",
            "If you could share one piece of wisdom with the world, what would it be?",
            "What does living a good life mean to you?",
        ]
    },
    "legacy_reflection": {
        "title": "Chapter 9: Legacy & Final Reflections",
        "questions": [
            "What are you most grateful for in your life?",
            "What do you want people to remember about you?",
            "If you could write a letter to your younger self, what would you say?",
            "What brings you peace and contentment now?",
            "How would you like your story to end?",
        ]
    }
}

# ============================================================================
# GEMINI CLIENT SETUP
# ============================================================================

client = genai.Client(api_key=GEMINI_API_KEY)

chat = client.chats.create(
    model="gemini-2.5-flash",
    config=types.GenerateContentConfig(
        system_instruction=(
            "You are a warm, empathetic biography interviewer named 'Story Keeper'. "
            "You are helping someone document their life story for future generations. "
            "When given instructions about which question to ask, acknowledge the user's "
            "previous answer with genuine interest and compassion, then naturally transition "
            "to the next question. Keep your responses conversational (2-4 sentences for acknowledgment). "
            "Never reveal you're following a script. Show that every story matters."
        )
    )
)

# ============================================================================
# SESSION MANAGEMENT (Save/Load Progress)
# ============================================================================


class InterviewSession:
    """Manages the interview session, saving progress and collected stories."""

    def __init__(self):
        self.collected_stories = {}  # {category: [(question, answer), ...]}
        self.current_category_index = 0
        self.current_question_index = 0
        self.categories = list(BIOGRAPHY_QUESTIONS.keys())
        self.interview_complete = False
        self.subject_name = ""
        self.start_date = datetime.datetime.now().isoformat()

    def save(self, filepath=SESSION_FILE):
        """Save session to JSON file."""
        data = {
            "subject_name": self.subject_name,
            "collected_stories": self.collected_stories,
            "current_category_index": self.current_category_index,
            "current_question_index": self.current_question_index,
            "interview_complete": self.interview_complete,
            "start_date": self.start_date,
            "last_saved": datetime.datetime.now().isoformat()
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"\n💾 Progress saved to {filepath}")

    def load(self, filepath=SESSION_FILE):
        """Load session from JSON file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            self.subject_name = data.get("subject_name", "")
            self.collected_stories = data.get("collected_stories", {})
            self.current_category_index = data.get("current_category_index", 0)
            self.current_question_index = data.get("current_question_index", 0)
            self.interview_complete = data.get("interview_complete", False)
            self.start_date = data.get(
                "start_date", datetime.datetime.now().isoformat())
            print(f"📂 Loaded previous session for {self.subject_name}")
            return True
        except FileNotFoundError:
            return False

    def add_response(self, category, question, answer):
        """Add a response to the collected stories."""
        if category not in self.collected_stories:
            self.collected_stories[category] = []
        self.collected_stories[category].append({
            "question": question,
            "answer": answer,
            "timestamp": datetime.datetime.now().isoformat()
        })

    def get_current_question(self):
        """Get the current question to ask."""
        if self.current_category_index >= len(self.categories):
            return None, None, None

        category = self.categories[self.current_category_index]
        category_data = BIOGRAPHY_QUESTIONS[category]
        questions = category_data["questions"]

        if self.current_question_index >= len(questions):
            return None, None, None

        return category, category_data["title"], questions[self.current_question_index]

    def advance_question(self):
        """Move to the next question."""
        category = self.categories[self.current_category_index]
        questions = BIOGRAPHY_QUESTIONS[category]["questions"]

        self.current_question_index += 1

        if self.current_question_index >= len(questions):
            self.current_question_index = 0
            self.current_category_index += 1

            if self.current_category_index >= len(self.categories):
                self.interview_complete = True
                return False  # No more questions

        return True

    def get_progress(self):
        """Get interview progress as percentage."""
        total_questions = sum(
            len(BIOGRAPHY_QUESTIONS[cat]["questions"]) for cat in self.categories)
        answered = sum(len(stories)
                       for stories in self.collected_stories.values())
        return (answered / total_questions) * 100 if total_questions > 0 else 0


# ============================================================================
# SPEECH RECOGNITION
# ============================================================================

def listen_and_transcribe():
    """Captures audio from the mic and converts it to text."""
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("\n🎙️ Listening... (Speak now, or type 'skip' to skip)")
        recognizer.adjust_for_ambient_noise(source, duration=1)
        try:
            # Longer for storytelling
            audio = recognizer.listen(
                source, timeout=10, phrase_time_limit=220)
            text = recognizer.recognize_google(audio)
            print(f"\n🗣️ You said: {text}")
            return text
        except sr.WaitTimeoutError:
            print("⏰ No speech detected. You can also type your response:")
            return input("✍️  Type here: ").strip() or None
        except (sr.UnknownValueError, sr.RequestError):
            print("🔇 Could not understand. Please type your response:")
            return input("✍️  Type here: ").strip() or None


def listen_with_typing_fallback():
    """Listen for voice or accept typed input."""
    print(
        "\n🎤 [Voice] or ⌨️  [Type] - Press Enter to start voice, or type your response:")
    typed = input(">>> ").strip()
    if typed:
        return typed
    return listen_and_transcribe()


# ============================================================================
# BIOGRAPHY DOCUMENT GENERATION
# ============================================================================

def generate_biography_document(session, output_path=BIOGRAPHY_OUTPUT):
    """Generate a formatted Word document biography from collected stories."""

    if not DOCX_AVAILABLE:
        print("❌ Cannot generate document: python-docx not installed")
        print("   Install with: pip install python-docx")
        # Fall back to text file
        return generate_biography_text(session)

    print("\n📝 Generating your biography document...")
    print("   This may take a few minutes as we craft each chapter...\n")

    doc = Document()

    # ---- Title Page ----
    title = doc.add_heading(f"The Life Story of {session.subject_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"\nA Personal Biography\n\nRecorded {datetime.datetime.now().strftime('%B %Y')}")
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_page_break()

    # ---- Table of Contents ----
    doc.add_heading("Table of Contents", level=1)
    for category in session.categories:
        if category in session.collected_stories and session.collected_stories[category]:
            chapter_title = BIOGRAPHY_QUESTIONS[category]["title"]
            doc.add_paragraph(chapter_title, style='List Number')

    doc.add_page_break()

    # ---- Generate Each Chapter ----
    for category in session.categories:
        if category not in session.collected_stories or not session.collected_stories[category]:
            continue

        chapter_data = BIOGRAPHY_QUESTIONS[category]
        chapter_title = chapter_data["title"]
        stories = session.collected_stories[category]

        print(f"   ✍️  Writing {chapter_title}...")

        # Add chapter heading
        doc.add_heading(chapter_title, level=1)

        # Generate narrative from Q&A using Gemini
        narrative = generate_chapter_narrative(
            session.subject_name, chapter_title, stories)

        # Add the narrative as paragraphs
        paragraphs = narrative.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(12)

        doc.add_page_break()

    # ---- Afterword ----
    doc.add_heading("Afterword", level=1)
    afterword = doc.add_paragraph(
        f"This biography was lovingly recorded through a series of conversations "
        f"with {session.subject_name}. Every story, memory, and reflection shared here "
        f"represents a piece of a life well-lived. May these words serve as a treasure "
        f"for future generations to know, understand, and remember the person behind these pages."
    )
    afterword.paragraph_format.first_line_indent = Inches(0.5)

    # Save the document
    doc.save(output_path)
    print(f"\n✅ Biography saved to: {output_path}")

    return output_path


def generate_chapter_narrative(subject_name, chapter_title, stories):
    """Use Gemini to transform Q&A into flowing narrative prose."""

    # Compile Q&A pairs
    qa_text = "\n\n".join([
        f"Question: {s['question']}\nAnswer: {s['answer']}"
        for s in stories
    ])

    prompt = f"""Transform the following interview Q&A about {subject_name}'s life into 
a compelling narrative chapter for their biography. 

Chapter: {chapter_title}

Interview Responses:
{qa_text}

Instructions:
- Write in third person, past tense
- Create flowing, engaging prose (not Q&A format)
- Add descriptive language and emotional depth
- Connect the stories naturally
- Aim for 2-3 pages of content (approximately 600-900 words)
- Make it read like a professional biography
- Preserve all the important details and quotes from the original responses

Write the narrative chapter:"""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        return response.text
    except Exception as e:
        print(f"⚠️  Error generating narrative: {e}")
        # Fallback: return formatted Q&A
        return "\n\n".join([f"{s['answer']}" for s in stories])


def generate_biography_text(session, output_path="my_biography.txt"):
    """Fallback: Generate a text file biography."""

    print("\n📝 Generating text biography...")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"THE LIFE STORY OF {session.subject_name.upper()}\n")
        f.write(f"{'='*50}\n\n")
        f.write(f"Recorded {datetime.datetime.now().strftime('%B %Y')}\n\n")

        for category in session.categories:
            if category not in session.collected_stories:
                continue

            chapter_title = BIOGRAPHY_QUESTIONS[category]["title"]
            stories = session.collected_stories[category]

            f.write(f"\n{chapter_title}\n")
            f.write(f"{'-'*40}\n\n")

            for story in stories:
                f.write(f"Q: {story['question']}\n")
                f.write(f"A: {story['answer']}\n\n")

    print(f"✅ Biography saved to: {output_path}")
    return output_path


# ============================================================================
# MAIN INTERVIEW APPLICATION
# ============================================================================

def run_interview():
    """Main interview loop."""

    print("\n" + "="*60)
    print("  📖 BIOGRAPHY GENERATOR - Voice-Enabled Life Story Interview")
    print("="*60)

    session = InterviewSession()

    # Check for existing session
    if Path(SESSION_FILE).exists():
        print("\n🔍 Found a previous interview session.")
        choice = input(
            "   Continue previous session? (yes/no): ").strip().lower()
        if choice in ['yes', 'y']:
            session.load()
            print(f"   📊 Progress: {session.get_progress():.1f}% complete")
        else:
            print("   Starting fresh interview...")

    # Get subject name if new session
    if not session.subject_name:
        print("\n👤 Who is this biography for?")
        session.subject_name = input("   Name: ").strip() or "Unknown"

    print(f"\n🎬 Starting biography interview for {session.subject_name}")
    print("-"*60)
    print("Commands during interview:")
    print("  • 'skip'  - Skip current question")
    print("  • 'back'  - Repeat current question")
    print("  • 'save'  - Save progress and continue")
    print("  • 'done'  - Finish and generate biography")
    print("  • 'quit'  - Save and exit (resume later)")
    print("-"*60)

    # Initial greeting
    category, chapter_title, first_question = session.get_current_question()
    if category:
        greeting = chat.send_message(
            f"Warmly greet {session.subject_name} and explain you're here to help capture their life story. "
            f"We're starting with {chapter_title}. Ask this question: '{first_question}'"
        )
        print(f"\n🤖 Story Keeper: {greeting.text}")

    # Main interview loop
    while not session.interview_complete:
        category, chapter_title, current_question = session.get_current_question()

        if not category:
            break

        # Get user response (voice or typed)
        user_text = listen_with_typing_fallback()

        if not user_text:
            continue

        # Handle commands
        cmd = user_text.lower().strip()

        if cmd in ['quit', 'exit']:
            session.save()
            print("\n👋 Interview paused. Your progress is saved!")
            print(f"   Run the program again to continue.")
            return

        if cmd == 'done':
            session.save()
            print("\n🎉 Finishing interview and generating biography...")
            break

        if cmd == 'save':
            session.save()
            print("   Continuing interview...")
            continue

        if cmd == 'skip':
            print("   ⏭️  Skipping question...")
            session.advance_question()

            # Ask next question
            next_cat, next_title, next_q = session.get_current_question()
            if next_q:
                if next_cat != category:
                    response = chat.send_message(
                        f"We're moving to a new chapter: {next_title}. "
                        f"Transition smoothly and ask: '{next_q}'"
                    )
                else:
                    response = chat.send_message(
                        f"Moving on. Ask this question: '{next_q}'")
                print(f"\n🤖 Story Keeper: {response.text}")
            continue

        if cmd == 'back':
            print(f"\n📋 Current question: {current_question}")
            continue

        # Record the response
        session.add_response(category, current_question, user_text)

        # Move to next question
        has_more = session.advance_question()

        if has_more:
            next_cat, next_title, next_q = session.get_current_question()

            if next_cat != category:
                # New chapter
                prompt = (
                    f"The user said: '{user_text}'. "
                    f"Acknowledge their response warmly. Then announce we're moving to "
                    f"the next chapter: {next_title}. Ask this question: '{next_q}'"
                )
            else:
                # Same chapter, next question
                prompt = (
                    f"The user said: '{user_text}'. "
                    f"Acknowledge thoughtfully, then ask: '{next_q}'"
                )

            print("\n🤖 Thinking...")
            response = chat.send_message(prompt)
            print(f"\n🤖 Story Keeper: {response.text}")
        else:
            # Interview complete
            final = chat.send_message(
                f"The user said: '{user_text}'. "
                f"This was the final question. Thank {session.subject_name} deeply for sharing "
                f"their incredible life story. Let them know their biography will now be created."
            )
            print(f"\n🤖 Story Keeper: {final.text}")

        # Auto-save every 5 responses
        total_responses = sum(len(s)
                              for s in session.collected_stories.values())
        if total_responses % 5 == 0:
            session.save()
            print(f"   📊 Progress: {session.get_progress():.1f}% complete")

    # Generate the biography
    session.interview_complete = True
    session.save()

    print("\n" + "="*60)
    print("  📚 GENERATING YOUR BIOGRAPHY")
    print("="*60)

    output_file = generate_biography_document(session)

    print("\n" + "🌟"*30)
    print(f"\n  ✅ Biography complete!")
    print(f"  📄 File: {output_file}")
    print(
        f"  📊 Total stories collected: {sum(len(s) for s in session.collected_stories.values())}")
    print("\n" + "🌟"*30)


# ============================================================================
# ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    print("\n" + "📖"*20)
    print("\n  BIOGRAPHY GENERATOR")
    print("  Capture Life Stories Through Conversation")
    print("\n" + "📖"*20)

    print("\nHow would you like to interact?")
    print("  1. Voice + Typing (recommended)")
    print("  2. Typing only")

    mode = input("\nChoice (1/2): ").strip()

    if mode == "2":
        # Override the listen function for typing only
        def listen_with_typing_fallback():
            return input("\n✍️  Your response: ").strip() or None

    try:
        run_interview()
    except KeyboardInterrupt:
        print("\n\n⏸️  Interview interrupted. Run again to continue.")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("   Your progress should be auto-saved. Run again to continue.")
